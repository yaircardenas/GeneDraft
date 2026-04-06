from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("GeneDraft_bioRxiv_preprint_draft.docx")


TITLE = "GeneDraft: a lightweight desktop workbench for DNA, RNA, and protein sequence editing and exploratory analysis"
AUTHOR = "Yair Cárdenas-Conejo*1,2"
AFFILIATIONS = [
    "1. Secretaría de Ciencia, Humanidades, Tecnología e Innovación, Gobierno de México, Ciudad de México, México",
    "2. Centro Universitario de Investigaciones Biomédicas, Universidad de Colima, Colima, Colima, México",
]
CORRESPONDING = (
    "Corresponding Author: Yair Cárdenas-Conejo* (ycardenasco@secihti.mx)\n"
    "ORCID: 0000-0002-0190-244X"
)

ABSTRACT = (
    "GeneDraft is a local desktop workbench for small-to-medium molecular sequence tasks "
    "that are often split across separate editors, web services, and ad hoc scripts. The "
    "software integrates sequence editing, FASTA and GenBank import/export, sequence "
    "validation, motif search, open reading frame (ORF) detection, selection-based "
    "translation, local DNA/RNA secondary-structure prediction, feature annotation, marker "
    "management, heuristic circularization review, SVG export, and BLAST workflows within a "
    "single graphical interface. In its current implementation, GeneDraft supports DNA, RNA, "
    "and protein sequences and combines local analysis with optional network-enabled "
    "functions including accession retrieval, web BLAST launching, and remote protein "
    "annotation enrichment. Protein-oriented reports include physicochemical summaries, "
    "transmembrane-segment inspection, heuristic expression-oriented interpretation, and "
    "optional enrichment through Protein-Sol, SoluProt, and Pfam/HMMER. GeneDraft is "
    "implemented in Python with a tkinter-based user interface, uses Biopython for sequence "
    "handling, and uses ViennaRNA for local thermodynamics-based secondary-structure "
    "prediction [12]. The software is intended for rapid exploratory work, construct review, "
    "teaching, and bench-adjacent sequence inspection rather than for genome-scale or "
    "high-throughput analysis. Its main contribution is not a new analysis algorithm but a "
    "local-first, editor-centered workflow that reduces context switching for routine "
    "sequence tasks."
)

KEYWORDS = (
    "sequence analysis; sequence editing; protein analysis; annotation; ORF detection; "
    "BLAST; FASTA; GenBank; desktop software; molecular biology"
)

SECTIONS = [
    (
        "Introduction",
        [
            "Routine molecular sequence work often consists of many small but essential "
            "operations: opening a FASTA or GenBank record, cleaning and validating sequence "
            "text, checking sequence type and composition, reviewing candidate motifs or open "
            "reading frames, translating a selected region, annotating features, and exporting "
            "the result for downstream analysis. Integrated desktop platforms have shown the "
            "value of bringing these tasks into a single environment [1]. At the same time, "
            "many everyday workflows remain distributed across separate editors, browser-based "
            "services, command-line tools, and one-off scripts. This fragmentation adds "
            "friction to exploratory work, complicates teaching demonstrations, and increases "
            "the chance of producing inconsistent intermediate files.",
            "Existing software covers many parts of this landscape. Mature integrated platforms "
            "such as Geneious and DNASTAR Lasergene/EditSeq are widely used in molecular "
            "biology [9,10], and open-source alternatives such as UGENE are available [11]. "
            "However, for users whose main need is rapid editor-centered inspection rather than "
            "maximal analytical breadth, even capable general-purpose environments may involve "
            "more interface overhead than is desirable for quick bench-adjacent tasks. In that "
            "context, the practical gap is often not the absence of algorithms but the absence "
            "of a lightweight local workspace that keeps frequent sequence operations close "
            "together.",
            "GeneDraft was developed to address that narrower but common use case. The goal of "
            "the software is not to replace specialized genome browsers, large-scale "
            "annotation systems, comparative genomics environments, or production "
            "bioinformatics pipelines. Instead, GeneDraft is intended as a local-first desktop "
            "workbench for rapid sequence inspection, lightweight annotation, exploratory "
            "translation and folding, quick protein review, and practical export in a single "
            "interface.",
            "The current version extends this editor-centered workflow beyond nucleotide "
            "handling by supporting protein-oriented review in the same environment. This is "
            "useful when users move directly from nucleotide inspection to translation-derived "
            "protein interpretation, or when they need a compact local report before deciding "
            "whether a sequence warrants more specialized downstream analysis.",
        ],
    ),
    (
        "Software architecture and implementation",
        [
            "GeneDraft is implemented in Python and uses tkinter for the desktop graphical "
            "interface. Core sequence parsing, translation, sequence summaries, GenBank feature "
            "handling, and file export are implemented with support from Biopython [2]. Local "
            "secondary-structure prediction for selected DNA or RNA regions is implemented "
            "through the ViennaRNA library [12]. The application follows a local-first design "
            "for its core functions: routine editing, validation, reverse-complement "
            "operations, DNA-to-RNA conversion, motif search, ORF detection, translation, "
            "secondary-structure prediction, feature editing, coordinate handling, "
            "circularization heuristics, and SVG map export do not require an internet "
            "connection.",
            "The software currently works with FASTA and GenBank inputs, also accepts sequences "
            "by direct copy-and-paste into the editor, and supports export back to both file "
            "formats. Sequence features imported from GenBank can be reviewed, focused, edited, "
            "deleted, or written back during GenBank export. The interface also supports named "
            "markers, recent-file tracking, persistent local settings, and recovery of unsaved "
            "sessions from the previous run. These functions aim to preserve continuity during "
            "exploratory analysis rather than enforce a formal project database model.",
            "Optional network-enabled functions are available for tasks that benefit from "
            "external resources. These include retrieval of records by NCBI accession, web BLAST "
            "launching, and remote enrichment of protein analysis through public online "
            "services. Local BLAST workflows are also supported when BLAST+ executables are "
            "installed in the user environment and a compatible database has been configured or "
            "built from FASTA input [3].",
        ],
    ),
    (
        "Functional scope",
        [
            "At its current stage, GeneDraft integrates file-based and editor-based sequence "
            "handling within a single desktop interface. Users can open FASTA and GenBank "
            "records, paste sequences directly into the editor, save curated content back to "
            "FASTA or GenBank, and work with DNA, RNA, or protein inputs while monitoring "
            "summary metrics such as sequence length, GC content, cursor position, and selection "
            "size. The same environment supports validation of invalid characters, "
            "normalization, reverse-complement operations, and interactive recovery of "
            "unfinished sessions.",
            "The analytical layer is organized around an editor-centered exploratory workflow. "
            "Within the same workspace, users can search motifs with strand awareness, detect "
            "ORFs with a configurable minimum size, translate selected regions across forward and "
            "reverse reading frames, convert DNA to RNA when needed, and fold selected DNA or "
            "RNA regions locally to obtain dot-bracket structure, minimum free energy, and an "
            "SVG secondary-structure diagram through ViennaRNA [12]. The same workspace "
            "supports creation and refinement of annotations through feature editing, coordinate "
            "export, marker placement, and GenBank feature import/export. The software also "
            "supports heuristic circularization analysis based on terminal overlap, allows "
            "circularization candidates to be marked or cropped, exports linear and circular "
            "maps as SVG, and provides both local and web-connected downstream paths through "
            "BLAST integration, local BLAST database creation, accession retrieval, and "
            "operation history for interactive editing.",
        ],
    ),
    (
        "Protein analysis workflow",
        [
            "One of the main additions in the current implementation is a dedicated "
            "protein-analysis workflow. This module accepts protein sequences directly or "
            "sequences obtained by translating nucleotide regions and generates a compact report "
            "intended for fast exploratory interpretation. The local analysis includes molecular "
            "weight, isoelectric point, instability index, GRAVY score, aromaticity, "
            "amino-acid composition, secondary-structure fractions, cysteine counts, and a "
            "simple transmembrane-segment scan based on a Kyte-Doolittle window [8].",
            "GeneDraft also derives heuristic expression-oriented summaries from these local "
            "properties. These include a local solubility estimate, a coarse crystallizability "
            "score, a simplified estimate of heterologous expression favorability in "
            "Escherichia coli, and an interpretation of likely soluble, membrane-associated, or "
            "aggregation-prone behavior. These outputs are intended as lightweight "
            "decision-support cues for exploratory use and not as substitutes for dedicated "
            "protein-expression or structural-biology pipelines.",
            "When network access is available, the protein-analysis report can be enriched "
            "asynchronously with remote predictors and domain annotation services. In the current "
            "version, these optional enrichments include Protein-Sol [4] and SoluProt [5] "
            "solubility scores and Pfam-domain detection through the EBI HMMER service [6], "
            "which builds on accelerated profile hidden Markov model search methods implemented "
            "in HMMER3 [7]. The local report is shown immediately, and remote annotations are "
            "inserted when available. This design keeps the workflow responsive while allowing "
            "best-effort access to external resources.",
        ],
    ),
    (
        "Positioning and intended use",
        [
            "GeneDraft is intended for rapid exploratory sequence work, especially in settings "
            "where users need to move quickly between inspection, annotation, translation, "
            "lightweight protein review, and export. Representative use cases include "
            "small-to-medium scale molecular biology projects, construct review, preliminary "
            "annotation of cloned or assembled regions, classroom demonstrations, and quick "
            "inspection of viral or plasmid-like components before downstream analysis in more "
            "specialized platforms.",
            "The software is particularly useful when the main need is to reduce transitions "
            "between independent tools rather than to maximize analytical depth in a single "
            "domain. Its value lies in integration, responsiveness, and low setup burden for "
            "common sequence tasks.",
        ],
    ),
    (
        "Limitations",
        [
            "GeneDraft is not designed as a replacement for genome-scale analysis environments, "
            "high-throughput annotation systems, formal comparative genomics workflows, or "
            "advanced structural-biology software. Circularization support is heuristic and "
            "based on terminal overlap rather than full alignment. Protein-expression and "
            "solubility outputs include simple local rules and optional web-derived annotations, "
            "so they should be interpreted as exploratory guidance rather than validated "
            "predictive endpoints. Likewise, local secondary-structure predictions are intended "
            "as rapid thermodynamics-based guidance for selected regions rather than as "
            "comprehensive structural modeling. Local BLAST requires BLAST+ command-line tools "
            "to be installed and accessible. Some optional functions depend on external services "
            "and therefore inherit the availability and interface stability of those services.",
            "At present, the software is designed for cross-platform desktop use but has been "
            "primarily tested on Windows. As with any exploratory sequence-analysis application, "
            "biological interpretation and downstream analytical decisions remain the "
            "responsibility of the user.",
        ],
    ),
    (
        "Availability and source code",
        [
            "Project name: GeneDraft",
            "Programming language: Python",
            "Main dependencies: tkinter, Biopython, ViennaRNA, optional NCBI BLAST+",
            "License: MIT License",
            "Version described here: 1.0",
        ],
    ),
    (
        "Conclusion",
        [
            "GeneDraft provides a compact local environment for DNA, RNA, and protein sequence "
            "editing, annotation, and exploratory analysis. By combining common sequence-handling "
            "operations with local secondary-structure prediction for selected nucleotide "
            "regions, optional accession retrieval, BLAST integration, and lightweight protein "
            "characterization, the software reduces workflow fragmentation during early-stage "
            "molecular analysis. GeneDraft is therefore best viewed as a practical desktop "
            "workbench for rapid sequence review and preparation, with particular value in "
            "research and teaching settings that benefit from fast iteration and minimal tool "
            "switching.",
        ],
    ),
    (
        "References",
        [],
    ),
    (
        "Declarations",
        [
            "Conflicts of interest: The author declares no conflicts of interest.",
            "Author contributions: Y.C.-C. conceived the project, developed the software, curated "
            "the documentation, and wrote the manuscript.",
        ],
    ),
]

CAPABILITIES = [
    "opening FASTA and GenBank files",
    "saving edited sequences to FASTA and GenBank",
    "sequence-type detection for DNA, RNA, and protein inputs",
    "sequence summaries including length, GC content, cursor position, and selection size",
    "validation and highlighting of invalid characters",
    "sequence normalization and reverse-complement operations",
    "motif search with forward and reverse-strand awareness",
    "ORF detection across the full sequence with configurable minimum size",
    "translation of selected regions across forward and reverse reading frames",
    "feature creation, editing, deletion, focusing, coordinate export, and GenBank feature import/export",
    "marker placement and coordinate tracking",
    "heuristic circularization analysis based on terminal overlap within a selected region",
    "marking or cropping circularization candidates",
    "export of linear and circular sequence maps as SVG",
    "BLAST web and local BLAST entry points from editor or results selections",
    "local BLAST database creation with makeblastdb",
    "retrieval of nucleotide or protein records by accession",
    "session recovery and operation history for interactive editing",
]

PENDING_ITEMS = [
    "Source code repository: https://github.com/yaircardenas/GeneDraft",
    "Archived release DOI: [Zenodo DOI pending]",
    "Operating system support: designed for Windows, Linux, and macOS; currently tested on Windows",
    "Funding: [Funding statement pending]",
    "Data and software availability: GeneDraft source code is available under the MIT License at "
    "https://github.com/yaircardenas/GeneDraft. Archived releases are available through Zenodo at [Zenodo DOI pending].",
]

REFERENCES = [
    "1. Kearse M, Moir R, Wilson A, Stones-Havas S, Cheung M, Sturrock S, et al. Geneious Basic: "
    "an integrated and extendable desktop software platform for the organization and analysis of "
    "sequence data. Bioinformatics. 2012;28(12):1647-1649. doi:10.1093/bioinformatics/bts199.",
    "2. Cock PJA, Antao T, Chang JT, Chapman BA, Cox CJ, Dalke A, et al. Biopython: freely "
    "available Python tools for computational molecular biology and bioinformatics. "
    "Bioinformatics. 2009;25(11):1422-1423. doi:10.1093/bioinformatics/btp163.",
    "3. Camacho C, Coulouris G, Avagyan V, Ma N, Papadopoulos J, Bealer K, et al. BLAST+: "
    "architecture and applications. BMC Bioinformatics. 2009;10:421. "
    "doi:10.1186/1471-2105-10-421.",
    "4. Hebditch M, Carballo-Amador MA, Charonis S, Curtis R, Warwicker J. Protein-Sol: a web "
    "tool for predicting protein solubility from sequence. Bioinformatics. "
    "2017;33(19):3098-3100. doi:10.1093/bioinformatics/btx345.",
    "5. Hon J, Marusiak M, Martinek T, Kunka A, Zendulka J, Bednar D, et al. SoluProt: prediction "
    "of soluble protein expression in Escherichia coli. Bioinformatics. 2021;37(1):23-28. "
    "doi:10.1093/bioinformatics/btaa1102.",
    "6. Potter SC, Luciani A, Eddy SR, Park Y, Lopez R, Finn RD. HMMER web server: 2018 update. "
    "Nucleic Acids Res. 2018;46(W1):W200-W204. doi:10.1093/nar/gky448.",
    "7. Eddy SR. Accelerated Profile HMM Searches. PLoS Comput Biol. 2011;7(10):e1002195. "
    "doi:10.1371/journal.pcbi.1002195.",
    "8. Kyte J, Doolittle RF. A simple method for displaying the hydropathic character of a "
    "protein. J Mol Biol. 1982;157(1):105-132. doi:10.1016/0022-2836(82)90515-0.",
    "9. Geneious Prime pricing. Geneious. https://www.geneious.com/pricing. Accessed April 6, "
    "2026.",
    "10. Lasergene Molecular Biology sequence analysis software. DNASTAR. "
    "https://www.dnastar.com/software/lasergene/molecular-biology/. Accessed April 6, 2026.",
    "11. Okonechnikov K, Golosova O, Fursov M, UGENE team. Unipro UGENE: a unified "
    "bioinformatics toolkit. Bioinformatics. 2012;28(8):1166-1167. "
    "doi:10.1093/bioinformatics/bts091.",
    "12. Lorenz R, Bernhart SH, Honer Zu Siederdissen C, Tafer H, Flamm C, Stadler PF, et al. "
    "ViennaRNA Package 2.0. Algorithms Mol Biol. 2011;6:26. doi:10.1186/1748-7188-6-26.",
]


def set_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    set_page_number(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"

    doc.styles["Title"].font.size = Pt(16)
    doc.styles["Title"].font.bold = True
    doc.styles["Heading 1"].font.size = Pt(12)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(12)
    doc.styles["Heading 2"].font.bold = True


def add_centered(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: int | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    if size is not None:
        run.font.size = Pt(size)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_before = Pt(10)
    p_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_pending(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True
    run.font.highlight_color = 7  # yellow
    run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)


def main() -> None:
    doc = Document()
    style_document(doc)

    add_centered(doc, TITLE, bold=True, size=16)
    add_centered(doc, AUTHOR, size=12)
    for item in AFFILIATIONS:
        add_centered(doc, item, size=11)
    add_centered(doc, CORRESPONDING, italic=True, size=11)

    add_heading(doc, "Abstract")
    add_paragraph(doc, ABSTRACT)

    add_heading(doc, "Keywords")
    add_paragraph(doc, KEYWORDS)

    for heading, paragraphs in SECTIONS:
        add_heading(doc, heading)
        for paragraph in paragraphs:
            add_paragraph(doc, paragraph)
        if heading == "References":
            for item in REFERENCES:
                add_paragraph(doc, item)
        if heading == "Availability and source code":
            for pending in PENDING_ITEMS[:3]:
                add_pending(doc, pending)
        if heading == "Declarations":
            for pending in PENDING_ITEMS[3:]:
                add_pending(doc, pending)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
